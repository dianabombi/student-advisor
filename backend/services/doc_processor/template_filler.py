"""
Template Filler

Fills document templates with extracted data.
"""

import logging
import os
import tempfile
from typing import Dict, Optional, List
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
from pathlib import Path

from .storage import MinIOStorage

logger = logging.getLogger(__name__)


class TemplateFiller:
    """
    Fill document templates with extracted data.
    
    Supports:
    - DOCX templates
    - Placeholder replacement
    - Table filling
    - Conditional sections
    """
    
    def __init__(self, storage: Optional[MinIOStorage] = None):
        """Initialize template filler with optional MinIO storage."""
        self.storage = storage or MinIOStorage()
        self.templates_dir = Path(__file__).parent / 'templates'
    
    def fill_template(
        self,
        template_path: str,
        data: Dict,
        output_path: str
    ) -> str:
        """
        Fill a DOCX template with data.
        
        Args:
            template_path: Path to template file
            data: Dictionary with field values
            output_path: Path for output file
            
        Returns:
            Path to filled document
        """
        logger.info(f"Filling template: {template_path}")
        
        # Load template
        doc = Document(template_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_placeholders(paragraph, data)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders(paragraph, data)
        
        # Save filled document
        doc.save(output_path)
        
        logger.info(f"Template filled: {output_path}")
        return output_path
    
    def _replace_placeholders(self, paragraph, data: Dict):
        """
        Replace placeholders in a paragraph.
        
        Args:
            paragraph: Document paragraph
            data: Data dictionary
        """
        text = paragraph.text
        
        # Replace each placeholder
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"  # {{key}} format
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        
        # Update paragraph text
        if text != paragraph.text:
            paragraph.text = text
    
    def create_contract_from_data(
        self,
        contract_type: str,
        data: Dict,
        output_path: str
    ) -> str:
        """
        Create a contract document from extracted data.
        
        Args:
            contract_type: Type of contract (e.g., 'kupna_zmluva')
            data: Extracted data
            output_path: Output path
            
        Returns:
            Path to created document
        """
        logger.info(f"Creating {contract_type} from data")
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(self._get_contract_title(contract_type), 0)
        
        # Add contract number and date
        if 'contract_number' in data.get('identifiers', {}):
            doc.add_paragraph(f"Číslo zmluvy: {data['identifiers']['contract_number']}")
        
        if data.get('dates'):
            doc.add_paragraph(f"Dátum: {data['dates'][0]['value']}")
        
        # Add parties
        doc.add_heading('Zmluvné strany', level=1)
        for i, party in enumerate(data.get('parties', []), 1):
            doc.add_paragraph(f"{i}. {party}")
        
        # Add amounts
        if data.get('amounts'):
            doc.add_heading('Finančné podmienky', level=1)
            for amount in data['amounts']:
                doc.add_paragraph(f"Suma: {amount['value']}")
                doc.add_paragraph(f"Kontext: {amount['context']}")
        
        # Add identifiers
        if data.get('identifiers'):
            doc.add_heading('Identifikácia', level=1)
            for key, value in data['identifiers'].items():
                doc.add_paragraph(f"{key.upper()}: {value}")
        
        # Save document
        doc.save(output_path)
        
        logger.info(f"Contract created: {output_path}")
        return output_path
    
    def _get_contract_title(self, contract_type: str) -> str:
        """Get human-readable contract title."""
        titles = {
            'kupna_zmluva': 'KÚPNA ZMLUVA',
            'najomna_zmluva': 'NÁJOMNÁ ZMLUVA',
            'zmluva_o_dielo': 'ZMLUVA O DIELO',
            'zmluva': 'ZMLUVA'
        }
        return titles.get(contract_type, 'DOKUMENT')
    
    def generate_summary(self, data: Dict) -> str:
        """
        Generate text summary of extracted data.
        
        Args:
            data: Extracted data
            
        Returns:
            Summary text
        """
        summary_parts = []
        
        # Document type
        if 'classification' in data:
            doc_type = data['classification'].get('document_type', 'unknown')
            summary_parts.append(f"Typ dokumentu: {doc_type}")
        
        # Parties
        if data.get('parties'):
            summary_parts.append(f"Strany: {', '.join(data['parties'][:3])}")
        
        # Dates
        if data.get('dates'):
            dates_str = ', '.join([d['value'] for d in data['dates'][:3]])
            summary_parts.append(f"Dátumy: {dates_str}")
        
        # Amounts
        if data.get('amounts'):
            amounts_str = ', '.join([a['value'] for a in data['amounts'][:3]])
            summary_parts.append(f"Sumy: {amounts_str}")
        
        # Identifiers
        if data.get('identifiers'):
            for key, value in data['identifiers'].items():
                summary_parts.append(f"{key.upper()}: {value}")
        
        return '\n'.join(summary_parts)
    
    def fill_and_save(
        self,
        template_name: str,
        data: Dict,
        output_filename: str,
        save_to_minio: bool = True
    ) -> Dict:
        """
        Fill template and optionally save to MinIO.
        
        Args:
            template_name: Name of template file
            data: Data dictionary
            output_filename: Output filename
            save_to_minio: Whether to save to MinIO
            
        Returns:
            Dictionary with file paths and URLs
        """
        logger.info(f"Filling template: {template_name}")
        
        # Get template path
        template_path = self._get_template_path(template_name)
        
        # Create temp output file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
        
        # Fill template
        self.fill_template(template_path, data, output_path)
        
        result = {
            'local_path': output_path,
            'filename': output_filename
        }
        
        # Save to MinIO if requested
        if save_to_minio:
            with open(output_path, 'rb') as f:
                file_data = f.read()
            
            # Upload to filled-docs bucket
            object_name = self.storage.upload_filled_document(
                file_data=file_data,
                filename=output_filename,
                template_name=template_name,
                source_document=data.get('_metadata', {}).get('source_document')
            )
            
            # Get presigned URL
            url = self.storage.get_presigned_url(
                self.storage.BUCKET_FILLED,
                object_name
            )
            
            result['minio_path'] = object_name
            result['url'] = url
            
            logger.info(f"Saved to MinIO: {object_name}")
        
        return result
    
    def generate_document(
        self,
        document_type: str,
        data: Dict,
        output_filename: str,
        save_to_minio: bool = True
    ) -> Dict:
        """
        Generate document from data without template.
        
        Args:
            document_type: Type of document
            data: Extracted data
            output_filename: Output filename
            save_to_minio: Whether to save to MinIO
            
        Returns:
            Dictionary with file info
        """
        logger.info(f"Generating {document_type} document")
        
        # Create temp output file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
        
        # Create document
        self.create_contract_from_data(document_type, data, output_path)
        
        result = {
            'local_path': output_path,
            'filename': output_filename
        }
        
        # Save to MinIO if requested
        if save_to_minio:
            with open(output_path, 'rb') as f:
                file_data = f.read()
            
            object_name = self.storage.upload_filled_document(
                file_data=file_data,
                filename=output_filename,
                template_name=f"generated_{document_type}"
            )
            
            url = self.storage.get_presigned_url(
                self.storage.BUCKET_FILLED,
                object_name
            )
            
            result['minio_path'] = object_name
            result['url'] = url
        
        return result
    
    def convert_to_pdf(self, docx_path: str, pdf_path: str) -> str:
        """
        Convert DOCX to PDF (requires LibreOffice or similar).
        
        Args:
            docx_path: Path to DOCX file
            pdf_path: Path for PDF output
            
        Returns:
            Path to PDF file
        """
        try:
            import subprocess
            
            # Try using LibreOffice
            subprocess.run([
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path),
                docx_path
            ], check=True)
            
            logger.info(f"Converted to PDF: {pdf_path}")
            return pdf_path
        
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            logger.error(f"PDF conversion failed: {e}")
            logger.warning("LibreOffice not available, returning DOCX")
            return docx_path
    
    def _get_template_path(self, template_name: str) -> str:
        """
        Get full path to template file.
        
        Args:
            template_name: Template filename
            
        Returns:
            Full path to template
        """
        # Check if template exists locally
        local_path = self.templates_dir / template_name
        if local_path.exists():
            return str(local_path)
        
        # Try to download from MinIO templates bucket
        try:
            template_data = self.storage.download_template(template_name)
            
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(template_data)
                return tmp.name
        
        except Exception as e:
            logger.error(f"Template not found: {template_name}")
            raise FileNotFoundError(f"Template not found: {template_name}")
    
    def list_templates(self) -> List[str]:
        """
        List available templates.
        
        Returns:
            List of template names
        """
        templates = []
        
        # Local templates
        if self.templates_dir.exists():
            templates.extend([
                f.name for f in self.templates_dir.glob('*.docx')
            ])
        
        # MinIO templates
        try:
            minio_templates = self.storage.list_files(
                self.storage.BUCKET_TEMPLATES
            )
            templates.extend([t['name'] for t in minio_templates])
        except Exception as e:
            logger.warning(f"Could not list MinIO templates: {e}")
        
        return list(set(templates))
    
    def upload_template(
        self,
        template_path: str,
        template_name: Optional[str] = None,
        template_type: str = "general"
    ) -> str:
        """
        Upload template to MinIO.
        
        Args:
            template_path: Path to template file
            template_name: Optional name (default: filename)
            template_type: Type of template
            
        Returns:
            Object name in MinIO
        """
        if not template_name:
            template_name = os.path.basename(template_path)
        
        with open(template_path, 'rb') as f:
            template_data = f.read()
        
        object_name = self.storage.upload_template(
            file_data=template_data,
            template_name=template_name,
            template_type=template_type
        )
        
        logger.info(f"Uploaded template: {object_name}")
        return object_name


# Convenience functions
def fill_template(
    template_name: str,
    data: Dict,
    output_filename: str,
    save_to_minio: bool = True
) -> Dict:
    """
    Fill template with data.
    
    Args:
        template_name: Template filename
        data: Data dictionary
        output_filename: Output filename
        save_to_minio: Save to MinIO
        
    Returns:
        Dictionary with file info
    """
    filler = TemplateFiller()
    return filler.fill_and_save(
        template_name=template_name,
        data=data,
        output_filename=output_filename,
        save_to_minio=save_to_minio
    )


def generate_document(
    document_type: str,
    data: Dict,
    output_filename: str,
    save_to_minio: bool = True
) -> Dict:
    """
    Generate document from data.
    
    Args:
        document_type: Document type
        data: Extracted data
        output_filename: Output filename
        save_to_minio: Save to MinIO
        
    Returns:
        Dictionary with file info
    """
    filler = TemplateFiller()
    return filler.generate_document(
        document_type=document_type,
        data=data,
        output_filename=output_filename,
        save_to_minio=save_to_minio
    )
